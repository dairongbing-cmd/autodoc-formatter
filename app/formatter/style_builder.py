from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.models.ir import Paragraph, Run
from app.models.rules import FormattingRules, HeadingStyle

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def hex_to_rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip("#")
    if len(hex_color) == 3:
        hex_color = "".join(c * 2 for c in hex_color)
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    return RGBColor(r, g, b)


def apply_paragraph_style(docx_para, rules: FormattingRules, ir_para: Paragraph | None = None) -> None:
    pf = docx_para.paragraph_format
    pr = rules.paragraph

    pf.line_spacing = pr.line_spacing
    pf.space_before = Pt(pr.space_before_pt)
    pf.space_after = Pt(pr.space_after_pt)

    alignment = ir_para.alignment if ir_para and ir_para.alignment else pr.alignment
    if alignment in ALIGN_MAP:
        pf.alignment = ALIGN_MAP[alignment]

    if pr.first_line_indent_cm is not None:
        pf.first_line_indent = Cm(pr.first_line_indent_cm)


def apply_heading_style(docx_para, rules: FormattingRules, level: int) -> None:
    heading_style = rules.headings.get(level)
    if heading_style is None:
        return

    pf = docx_para.paragraph_format
    if heading_style.space_before_pt is not None:
        pf.space_before = Pt(heading_style.space_before_pt)
    if heading_style.space_after_pt is not None:
        pf.space_after = Pt(heading_style.space_after_pt)
    if heading_style.alignment and heading_style.alignment in ALIGN_MAP:
        pf.alignment = ALIGN_MAP[heading_style.alignment]

    for run in docx_para.runs:
        _apply_run_style(run, rules, heading_style)


def apply_run_style(docx_run, rules: FormattingRules, ir_run: Run | None = None) -> None:
    _apply_run_style(docx_run, rules, heading_style=None, ir_run=ir_run)


def _apply_run_style(
    docx_run,
    rules: FormattingRules,
    heading_style: HeadingStyle | None = None,
    ir_run: Run | None = None,
) -> None:
    if heading_style:
        if heading_style.font_family:
            docx_run.font.name = heading_style.font_family
        if heading_style.font_size_pt is not None:
            docx_run.font.size = Pt(heading_style.font_size_pt)
        if heading_style.color_hex:
            docx_run.font.color.rgb = hex_to_rgb(heading_style.color_hex)
        docx_run.bold = heading_style.bold
        docx_run.italic = heading_style.italic
    else:
        font_family = (ir_run.font_name if ir_run and ir_run.font_name else rules.font.family)
        docx_run.font.name = font_family

        font_size = (ir_run.font_size_pt if ir_run and ir_run.font_size_pt is not None else rules.font.size_pt)
        docx_run.font.size = Pt(font_size)

        if ir_run and ir_run.font_color_hex:
            docx_run.font.color.rgb = hex_to_rgb(ir_run.font_color_hex)
        else:
            docx_run.font.color.rgb = hex_to_rgb(rules.font.color_hex)

        docx_run.bold = ir_run.bold if ir_run else rules.font.bold_default
        docx_run.italic = ir_run.italic if ir_run else rules.font.italic_default
        docx_run.underline = ir_run.underline if ir_run else False

    if ir_run and ir_run.strikethrough:
        rpr = docx_run._r.get_or_add_rPr()
        strike = rpr.find(qn("w:strike"))
        if strike is None:
            from lxml import etree
            strike = etree.SubElement(rpr, qn("w:strike"))
        strike.set(qn("w:val"), "true")


def setup_page(docx_section, rules: FormattingRules) -> None:
    pr = rules.page
    page_map = {"A4": (Cm(21.0), Cm(29.7)), "Letter": (Cm(21.59), Cm(27.94))}
    width, height = page_map.get(pr.size, (Cm(21.0), Cm(29.7)))

    if pr.orientation == "landscape":
        width, height = height, width

    docx_section.page_width = width
    docx_section.page_height = height

    m = pr.margins_cm
    docx_section.top_margin = Cm(m.top_cm)
    docx_section.bottom_margin = Cm(m.bottom_cm)
    docx_section.left_margin = Cm(m.left_cm)
    docx_section.right_margin = Cm(m.right_cm)


def build_header_footer(docx_section, rules: FormattingRules) -> None:
    if rules.header.enabled:
        header = docx_section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        text = rules.header.text or ""
        if rules.header.show_page_numbers:
            text += "  "
            run = p.add_run()
            fld_char_begin = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
            run._r.append(fld_char_begin)
            run2 = p.add_run()
            instr = run2._r.makeelement(qn("w:instrText"), {})
            instr.text = " PAGE "
            run2._r.append(instr)
            run3 = p.add_run()
            fld_char_end = run3._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
            run3._r.append(fld_char_end)
        if text.strip():
            p.add_run(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if rules.footer.enabled:
        footer = docx_section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        text = rules.footer.text or ""
        if rules.footer.show_page_numbers:
            run = p.add_run()
            fld_char_begin = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
            run._r.append(fld_char_begin)
            run2 = p.add_run()
            instr = run2._r.makeelement(qn("w:instrText"), {})
            instr.text = " PAGE "
            run2._r.append(instr)
            run3 = p.add_run()
            fld_char_end = run3._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
            run3._r.append(fld_char_end)
        if text.strip():
            p.add_run(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
