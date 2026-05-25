from io import BytesIO

from docx import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, Pt

from app.models.ir import (
    Document,
    HorizontalRule,
    ImageReference,
    PageBreak,
    Paragraph,
    Run,
    Section,
    Table,
)
from app.models.rules import FormattingRules
from app.formatter.style_builder import (
    apply_heading_style,
    apply_paragraph_style,
    apply_run_style,
    build_header_footer,
    hex_to_rgb,
    setup_page,
)


class FormatEngine:
    def __init__(self, rules: FormattingRules):
        self.rules = rules

    def apply(self, document: Document) -> bytes:
        docx = DocxDocument()

        for i, section in enumerate(document.sections):
            if i > 0:
                docx.add_section()
            self._process_section(docx, section, document)

        # Apply page setup to all sections
        for section in docx.sections:
            setup_page(section, self.rules)
            build_header_footer(section, self.rules)

        buf = BytesIO()
        docx.save(buf)
        buf.seek(0)
        return buf.read()

    def _process_section(self, docx: DocxDocument, section: Section, document: Document) -> None:
        for block in section.blocks:
            if isinstance(block, Paragraph):
                self._write_paragraph(docx, block)
            elif isinstance(block, Table):
                self._write_table(docx, block)
            elif isinstance(block, ImageReference):
                self._write_image(docx, block, document)
            elif isinstance(block, PageBreak):
                docx.add_page_break()
            elif isinstance(block, HorizontalRule):
                p = docx.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                pPr = p._p.get_or_add_pPr()
                pBdr = pPr.makeelement(docx.oxml.qn("w:pBdr"), {})
                bottom = pBdr.makeelement(docx.oxml.qn("w:bottom"), {
                    docx.oxml.qn("w:val"): "single",
                    docx.oxml.qn("w:sz"): "12",
                    docx.oxml.qn("w:space"): "1",
                    docx.oxml.qn("w:color"): "auto",
                })
                pBdr.append(bottom)
                pPr.append(pBdr)

    def _write_paragraph(self, docx: DocxDocument, ir_para: Paragraph) -> None:
        if ir_para.style_hint and ir_para.style_hint.startswith("heading"):
            level = int(ir_para.style_hint[-1])
            text = "".join(r.text for r in ir_para.runs)
            docx_para = docx.add_heading(text, level=level)
            apply_heading_style(docx_para, self.rules, level)
        else:
            docx_para = docx.add_paragraph()
            apply_paragraph_style(docx_para, self.rules, ir_para)

            if ir_para.style_hint == "code":
                for ir_run in ir_para.runs:
                    docx_run = docx_para.add_run(ir_run.text)
                    docx_run.font.name = "Courier New"
                    docx_run.font.size = Pt(10)
            elif ir_para.style_hint == "quote":
                docx_para.paragraph_format.left_indent = Cm(1.5)
                for ir_run in ir_para.runs:
                    docx_run = docx_para.add_run(ir_run.text)
                    apply_run_style(docx_run, self.rules, ir_run)
                    docx_run.italic = True
            else:
                for ir_run in ir_para.runs:
                    docx_run = docx_para.add_run(ir_run.text)
                    apply_run_style(docx_run, self.rules, ir_run)

                if ir_para.list_level > 0:
                    if ir_para.list_type == "number":
                        docx_para.style = docx_para.part.styles["List Number"]
                    else:
                        docx_para.style = docx_para.part.styles["List Bullet"]

    def _write_table(self, docx: DocxDocument, ir_table: Table) -> None:
        if not ir_table.rows:
            return

        num_cols = max(len(row.cells) for row in ir_table.rows)
        if num_cols == 0:
            return

        docx_table = docx.add_table(rows=len(ir_table.rows), cols=num_cols)
        docx_table.style = "Table Grid"

        for r_idx, ir_row in enumerate(ir_table.rows):
            for c_idx, ir_cell in enumerate(ir_row.cells):
                if c_idx >= num_cols:
                    break
                cell = docx_table.cell(r_idx, c_idx)
                for block in ir_cell.blocks:
                    if isinstance(block, Paragraph):
                        text = "".join(r.text for r in block.runs)
                        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        p.clear()
                        p.add_run(text)
                        apply_paragraph_style(p, self.rules, block)

        if ir_table.caption:
            caption_para = docx.add_paragraph()
            caption_run = caption_para.add_run(ir_table.caption)
            caption_run.italic = True
            caption_run.font.size = Pt(10)
            caption_para.paragraph_format.alignment = docx.oxml.text.WD_ALIGN_PARAGRAPH.CENTER

    def _write_image(self, docx: DocxDocument, ir_image: ImageReference, document: Document) -> None:
        inline_img = None
        for img in document.images:
            if img.image_id == ir_image.image_id:
                inline_img = img
                break

        if inline_img is None:
            p = docx.add_paragraph()
            p.add_run(f"[图片: {ir_image.caption or ir_image.image_id}]")
            return

        try:
            from PIL import Image as PILImage
            img = PILImage.open(BytesIO(inline_img.data))
            width_px, height_px = img.size
        except Exception:
            width_px, height_px = None, None

        p = docx.add_paragraph()
        p.alignment = 1  # center
        run = p.add_run()

        max_width = self.rules.images.max_width_cm
        if max_width and width_px and height_px:
            aspect = height_px / width_px
            w_cm = min(max_width, width_px / 37.8)  # approximate px->cm conversion
            run.add_picture(BytesIO(inline_img.data), width=Cm(w_cm))
        else:
            run.add_picture(BytesIO(inline_img.data))

        if self.rules.images.add_captions and ir_image.caption:
            cap = docx.add_paragraph()
            cap.alignment = 1
            cap_run = cap.add_run(ir_image.caption)
            cap_run.italic = True
            cap_run.font.size = Pt(10)
